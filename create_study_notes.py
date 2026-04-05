"""
生成学习笔记 Word 文档
每完成一步就运行这个脚本更新笔记
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI AV Schematic Generator")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("从零搭建 AI 自动画原理图系统\n学习笔记")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f"开始日期: {datetime.now().strftime('%Y-%m-%d')}")
run.font.size = Pt(12)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading("项目总览", level=1)
doc.add_paragraph(
    "本项目目标：搭建一个 AI 驱动的 AV 原理图自动生成系统，\n"
    "同时学习 CI/CD、Git、部署等开发技能。"
)

doc.add_heading("学习路线图", level=2)
stages = [
    ("阶段 1", "搭建 AI 画图核心", "Python + Claude API + Graphviz"),
    ("阶段 2", "加 Web 界面", "Flask/FastAPI + HTML"),
    ("阶段 3", "放到 GitHub", "Git 基础操作"),
    ("阶段 4", "搭 CI/CD", "GitHub Actions 自动测试+部署"),
    ("阶段 5", "部署上线", "云服务器 / Vercel"),
]

table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
hdr = table.rows[0].cells
hdr[0].text = "阶段"
hdr[1].text = "内容"
hdr[2].text = "技术栈"
for stage, content, tech in stages:
    row = table.add_row().cells
    row[0].text = stage
    row[1].text = content
    row[2].text = tech

doc.add_page_break()

# ===== 阶段 1 =====
doc.add_heading("阶段 1：搭建 AI 画图核心", level=1)

# --- Step 1.1 ---
doc.add_heading("Step 1.1：环境检查", level=2)
doc.add_paragraph("目的：确认电脑上已安装必要工具。")

doc.add_heading("什么是这些工具？", level=3)
tools_explanation = [
    ("Python", "编程语言，我们写代码用的主要语言"),
    ("pip", "Python 的包管理器，用来安装第三方库（类似 App Store）"),
    ("Git", "版本控制工具，记录代码的每次修改（类似游戏存档）"),
    ("Node.js", "JavaScript 运行环境（后面做 Web 界面会用到）"),
    ("VS Code", "代码编辑器（写代码的地方）"),
]
for tool, explanation in tools_explanation:
    p = doc.add_paragraph()
    run = p.add_run(f"{tool}：")
    run.bold = True
    p.add_run(explanation)

doc.add_heading("检查命令", level=3)
doc.add_paragraph("在终端（Terminal）里输入以下命令：", style="List Bullet")

commands = [
    ("python --version", "查看 Python 版本"),
    ("pip --version", "查看 pip 版本"),
    ("git --version", "查看 Git 版本"),
    ("node --version", "查看 Node.js 版本"),
]
for cmd, desc in commands:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(cmd)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.add_run(f"  — {desc}")

doc.add_heading("你的检查结果", level=3)
results = [
    ("Python", "3.11.9", "OK"),
    ("Git", "2.53.0", "OK"),
    ("Node.js", "22.21.0", "OK"),
    ("pip", "26.0", "OK"),
    ("VS Code", "1.113.0", "OK"),
    ("Docker", "未安装", "后面再装"),
    ("GitHub CLI", "未安装", "后面再装"),
]
table2 = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
hdr2 = table2.rows[0].cells
hdr2[0].text = "工具"
hdr2[1].text = "版本"
hdr2[2].text = "状态"
for tool, ver, status in results:
    row = table2.add_row().cells
    row[0].text = tool
    row[1].text = ver
    row[2].text = status

doc.add_paragraph()
doc.add_paragraph("Step 1.1 完成！").bold = True

doc.add_page_break()

# --- Step 1.2 ---
doc.add_heading("Step 1.2：创建项目 & 安装依赖", level=2)
doc.add_paragraph("目的：创建项目文件夹，安装需要用到的 Python 库。")

doc.add_heading("关键概念", level=3)
concepts = [
    ("项目文件夹", "所有代码放在一个文件夹里，方便管理"),
    ("依赖/库 (Library)", "别人写好的代码，我们直接拿来用，不用重新造轮子"),
    ("pip install", "安装 Python 库的命令"),
    ("requirements.txt", "记录项目需要哪些库（类似购物清单）"),
]
for concept, explanation in concepts:
    p = doc.add_paragraph()
    run = p.add_run(f"{concept}：")
    run.bold = True
    p.add_run(explanation)

doc.add_heading("我们要安装的库", level=3)
libs = [
    ("anthropic", "Anthropic 官方 SDK，用来调用 Claude AI 的 API"),
    ("graphviz", "画图库，把文字描述变成图片"),
    ("python-dotenv", "读取 .env 文件，安全存储 API 密钥"),
]
for lib, desc in libs:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(lib)
    run.bold = True
    p.add_run(f" — {desc}")

doc.add_heading("操作步骤", level=3)
steps = [
    "打开终端 (Terminal)",
    "创建项目文件夹：mkdir ai-schematic-generator",
    "进入文件夹：cd ai-schematic-generator",
    "安装依赖：pip install anthropic graphviz python-dotenv",
    "创建 requirements.txt 记录依赖",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("什么是 API Key？", level=3)
doc.add_paragraph(
    "API Key 就像一把钥匙，让你的程序可以访问 Claude AI。\n"
    "每次你的程序发请求给 Claude，都需要带上这把钥匙证明身份。\n\n"
    "重要：API Key 绝对不能公开！不能放到 GitHub 上！\n"
    "所以我们用 .env 文件来存储，并且用 .gitignore 忽略它。"
)

doc.add_heading("如何获取 Claude API Key", level=3)
api_steps = [
    "打开 https://console.anthropic.com/",
    "注册/登录账号",
    "点击左边菜单 'API Keys'",
    "点击 'Create Key'",
    "复制生成的 Key（只显示一次！）",
    "粘贴到项目的 .env 文件里",
]
for i, step in enumerate(api_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("等你拿到 API Key，我们就进入 Step 1.3：写核心代码！")
run.bold = True

doc.add_page_break()

# --- Step 1.3 ---
doc.add_heading("Step 1.3：核心代码 - schematic_engine.py", level=2)
doc.add_paragraph("目的：写出系统的核心代码，实现 用户需求 → AI 分析 → 自动画图。")

doc.add_heading("整体流程", level=3)
doc.add_paragraph(
    "1. 用户输入 AV 系统需求（文字描述）\n"
    "2. 程序把需求发给 Claude AI\n"
    "3. Claude AI 返回 JSON 格式的设备清单和连接关系\n"
    "4. 程序用 Graphviz 把 JSON 画成原理图\n"
    "5. 输出 PNG 图片"
)

doc.add_heading("关键概念", level=3)
concepts_13 = [
    ("API (Application Programming Interface)",
     "程序之间沟通的接口。我们的程序通过 API 跟 Claude AI 对话。"),
    ("JSON (JavaScript Object Notation)",
     "一种数据格式，类似一张表格。AI 返回的数据用 JSON 格式，方便程序读取。"),
    ("Graphviz",
     "开源画图工具。我们告诉它'有哪些方块、怎么连线'，它自动排版画图。"),
    ("System Prompt",
     "给 AI 的'角色设定'。告诉 AI 它是 AV 工程师，要按特定格式回答。"),
    ("Demo Mode（模拟模式）",
     "没有 API Key 时，用预设数据运行，方便开发和测试。"),
    (".env 文件",
     "存储敏感信息（如 API Key）的文件，不会上传到 GitHub。"),
]
for concept, explanation in concepts_13:
    p = doc.add_paragraph()
    run = p.add_run(f"{concept}：")
    run.bold = True
    p.add_run(f"\n{explanation}")

doc.add_heading("代码结构解读", level=3)
doc.add_paragraph(
    "schematic_engine.py 文件分为 4 个部分：\n\n"
    "1. SYSTEM_PROMPT（系统提示词）\n"
    "   - 告诉 AI 它的角色是 AV 工程师\n"
    "   - 规定 AI 必须返回 JSON 格式\n"
    "   - 定义 JSON 的结构（devices + connections）\n\n"
    "2. DEMO_RESULT（模拟数据）\n"
    "   - 没有 API Key 时用的假数据\n"
    "   - 包含会议室 AV 系统的完整示例\n\n"
    "3. call_claude_ai() 函数\n"
    "   - 检查有没有 API Key\n"
    "   - 没有 → 返回模拟数据\n"
    "   - 有 → 调用 Claude API，获取真实回答\n\n"
    "4. generate_diagram() 函数\n"
    "   - 接收 JSON 数据\n"
    "   - 用 Graphviz 创建图：\n"
    "     * 蓝色方块 = INPUT 设备\n"
    "     * 橙色方块 = PROCESSING 设备\n"
    "     * 绿色方块 = OUTPUT 设备\n"
    "   - 箭头上标注线材类型和信号类型\n"
    "   - 输出 PNG 图片到 output/ 文件夹"
)

doc.add_heading("运行命令", level=3)
p = doc.add_paragraph()
run = p.add_run("python schematic_engine.py")
run.font.name = "Consolas"
run.font.size = Pt(10)

doc.add_heading("输出结果", level=3)
doc.add_paragraph(
    "运行后会在 output/ 文件夹生成 schematic.png\n"
    "这就是自动画出来的 AV 原理图！\n\n"
    "图中包含：\n"
    "- INPUT 区域（蓝色）：Laptop, Wireless Mic, PTZ Camera\n"
    "- PROCESSING 区域（橙色）：HDMI Switcher, QSYS Core, Amplifier, Zoom Codec\n"
    "- OUTPUT 区域（绿色）：85\" Display, Ceiling Speakers\n"
    "- 连接线标注：HDMI(Video), Dante/CAT6(Audio), XLR(Audio) 等"
)

doc.add_heading("切换到真实 AI 模式", level=3)
doc.add_paragraph(
    "当你拿到 Anthropic API Key 后：\n"
    "1. 打开 .env 文件\n"
    "2. 把 your-api-key-here 替换成你的真实 Key\n"
    "3. 重新运行 python schematic_engine.py\n"
    "4. 这时 AI 会根据你的输入实时生成原理图！"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Step 1.3 完成！阶段 1 全部完成！")
run.bold = True

doc.add_page_break()

# ===== 阶段 2 =====
doc.add_heading("阶段 2：Web 界面", level=1)

doc.add_heading("Step 2.1：什么是 Web 应用？", level=2)
doc.add_paragraph(
    "之前我们在终端（黑窗口）里运行程序，只有你自己能用。\n"
    "加上 Web 界面后，打开浏览器就能用，以后还能给客户用。\n\n"
    "Web 应用分两部分：\n"
    "- 前端 (Frontend)：用户看到的网页界面（HTML + CSS + JavaScript）\n"
    "- 后端 (Backend)：处理请求的服务器程序（Python + Flask）"
)

doc.add_heading("关键概念", level=3)
web_concepts = [
    ("Flask", "Python 的 Web 框架，帮你快速搭建网站后端。类似一个'骨架'，你往里面填内容就行。"),
    ("路由 (Route)", "URL 地址和功能的对应关系。比如访问 / 显示首页，访问 /generate 生成图。"),
    ("HTML", "网页的骨架语言，定义页面有哪些元素（标题、按钮、输入框等）。"),
    ("CSS", "网页的样式语言，定义元素长什么样（颜色、大小、位置等）。"),
    ("JavaScript", "网页的动作语言，定义用户操作时发生什么（点按钮→发请求→显示结果）。"),
    ("API 接口", "前端和后端之间的沟通方式。前端发 JSON 请求，后端返回 JSON 结果。"),
    ("JSON", "数据交换格式，像一个结构化的'信封'，前后端都能读懂。"),
    ("localhost:5000", "本地服务器地址。localhost = 你自己的电脑，5000 = 端口号（门牌号）。"),
]
for concept, explanation in web_concepts:
    p = doc.add_paragraph()
    run = p.add_run(f"{concept}：")
    run.bold = True
    p.add_run(f"\n{explanation}")

doc.add_heading("Step 2.2：项目文件结构", level=2)
doc.add_paragraph(
    "ai-schematic-generator/\n"
    "  app.py                  -- Web 服务器（后端）\n"
    "  schematic_engine.py     -- AI 核心引擎（阶段1写的）\n"
    "  requirements.txt        -- 依赖清单\n"
    "  .env                    -- API Key\n"
    "  templates/\n"
    "    index.html            -- 网页界面（前端）\n"
    "  static/                 -- 静态文件（图片、CSS等）\n"
    "  output/                 -- 生成的原理图"
)

doc.add_heading("Step 2.3：后端代码解读 (app.py)", level=2)
doc.add_paragraph(
    "app.py 做了 3 件事：\n\n"
    "1. @app.route('/')  --  首页路由\n"
    "   用户访问 http://localhost:5000 时，显示 index.html 页面\n\n"
    "2. @app.route('/generate')  --  生成路由\n"
    "   接收用户的 AV 需求（JSON 格式）\n"
    "   调用 schematic_engine.py 里的函数\n"
    "   返回设备清单 + 图片路径\n\n"
    "3. @app.route('/image/<filename>')  --  图片路由\n"
    "   提供生成的 PNG 图片给浏览器显示"
)

doc.add_heading("Step 2.4：前端代码解读 (index.html)", level=2)
doc.add_paragraph(
    "index.html 是一个单页面应用，包含：\n\n"
    "HTML 部分（骨架）：\n"
    "- 顶部导航栏：显示标题和模式标识\n"
    "- 输入区域：文本框 + 示例按钮 + 生成按钮\n"
    "- 结果区域：统计卡片 + 原理图 + 设备清单\n\n"
    "CSS 部分（样式）：\n"
    "- 深色主题设计（专业感）\n"
    "- 响应式布局（手机也能用）\n"
    "- 动画效果（加载中旋转图标）\n\n"
    "JavaScript 部分（动作）：\n"
    "- generate() 函数：点击按钮时发送请求到后端\n"
    "- showResult() 函数：收到结果后更新页面\n"
    "- fillExample() 函数：点击示例按钮填入预设需求\n"
    "- Ctrl+Enter 快捷键提交"
)

doc.add_heading("Step 2.5：如何运行", level=2)
doc.add_paragraph(
    "1. 打开终端，进入项目文件夹\n"
    "   cd ai-schematic-generator\n\n"
    "2. 启动 Web 服务器\n"
    "   python app.py\n\n"
    "3. 打开浏览器，访问\n"
    "   http://localhost:5000\n\n"
    "4. 在文本框输入 AV 需求（或点击示例按钮）\n\n"
    "5. 点击 Generate Schematic\n\n"
    "6. 等待几秒，原理图就出来了！\n\n"
    "7. 点击 Download PNG 下载图片\n\n"
    "8. 要停止服务器：在终端按 Ctrl+C"
)

doc.add_heading("前后端交互流程", level=3)
doc.add_paragraph(
    "用户点击按钮\n"
    "    |\n"
    "    v\n"
    "JavaScript 发送 POST 请求到 /generate\n"
    "（带着用户输入的文字）\n"
    "    |\n"
    "    v\n"
    "Flask 后端接收请求\n"
    "    |\n"
    "    v\n"
    "调用 call_claude_ai() 获取 AI 结果\n"
    "    |\n"
    "    v\n"
    "调用 generate_diagram() 画图\n"
    "    |\n"
    "    v\n"
    "返回 JSON（设备清单 + 图片路径）\n"
    "    |\n"
    "    v\n"
    "JavaScript 更新页面显示结果"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("阶段 2 完成！你现在有了一个可以在浏览器里使用的 AI 画图系统！")
run.bold = True

doc.add_page_break()

# ===== 阶段 3 =====
doc.add_heading("阶段 3：Git + GitHub", level=1)

doc.add_heading("Step 3.1：什么是 Git？", level=2)
doc.add_paragraph(
    "Git 是版本控制工具，帮你记录代码的每次修改。\n\n"
    "类比：\n"
    "- 没有 Git：写错了只能 Ctrl+Z，关掉文件就没了\n"
    "- 有 Git：每次'存档'都保留，随时可以回到任何版本\n\n"
    "核心概念：\n"
    "- Repository (仓库)：项目的文件夹，Git 会追踪里面的所有变化\n"
    "- Commit (提交)：一次'存档'，记录了你改了什么、为什么改\n"
    "- Stage (暂存)：选择哪些修改要放进下一次 commit\n"
    "- Branch (分支)：代码的平行世界，可以同时开发多个功能"
)

doc.add_heading("Step 3.2：什么是 GitHub？", level=2)
doc.add_paragraph(
    "GitHub 是代码托管平台（Git 的云端版本）。\n\n"
    "类比：\n"
    "- Git = 本地的游戏存档\n"
    "- GitHub = 把存档上传到云端，换电脑也能继续玩\n\n"
    "GitHub 的作用：\n"
    "- 备份代码（电脑坏了也不怕）\n"
    "- 多人协作（团队一起写代码）\n"
    "- 展示作品（面试、接单都能用）\n"
    "- CI/CD（自动测试和部署，阶段4会学）"
)

doc.add_heading("Step 3.3：什么是 .gitignore？", level=2)
doc.add_paragraph(
    ".gitignore 文件告诉 Git 哪些文件不要追踪。\n\n"
    "我们忽略了：\n"
    "- .env：API Key 绝对不能上传！\n"
    "- __pycache__/：Python 缓存，自动生成的\n"
    "- output/：生成的图片，每次运行会重新生成\n"
    "- *.docx：个人学习笔记\n\n"
    "重要安全原则：\n"
    "密码、API Key、私钥等敏感信息永远不要上传到 GitHub！\n"
    "用 .env 文件存储，用 .gitignore 忽略。"
)

doc.add_heading("Step 3.4：Git 常用命令", level=2)
git_commands = [
    ("git init", "初始化仓库（只需要做一次）"),
    ("git status", "查看当前状态（哪些文件被修改了）"),
    ("git add <file>", "把文件加入暂存区（选择要提交的文件）"),
    ("git commit -m 'message'", "提交（存档），message 描述你做了什么"),
    ("git log", "查看提交历史"),
    ("git push", "把本地代码推送到 GitHub"),
    ("git pull", "从 GitHub 拉取最新代码"),
    ("git diff", "查看修改了什么内容"),
]
table3 = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
hdr3 = table3.rows[0].cells
hdr3[0].text = "命令"
hdr3[1].text = "作用"
for cmd, desc in git_commands:
    row = table3.add_row().cells
    row[0].text = cmd
    row[1].text = desc

doc.add_heading("Step 3.5：GitHub CLI (gh)", level=2)
doc.add_paragraph(
    "GitHub CLI 是 GitHub 的命令行工具，可以在终端里操作 GitHub。\n\n"
    "安装命令：\n"
    "  winget install --id GitHub.cli\n\n"
    "登录命令：\n"
    "  gh auth login --web -p https\n\n"
    "创建远程仓库并推送：\n"
    "  gh repo create ai-schematic-generator --public --source=. --remote=origin --push\n\n"
    "这一条命令做了 4 件事：\n"
    "1. 在 GitHub 上创建新仓库\n"
    "2. --public 表示公开仓库（别人可以看到）\n"
    "3. --source=. 表示用当前文件夹的代码\n"
    "4. --push 表示立即推送代码上去"
)

doc.add_heading("Step 3.6：我们的仓库", level=2)
doc.add_paragraph(
    "仓库地址：https://github.com/yeahhao1997/ai-schematic-generator\n\n"
    "上传的文件：\n"
    "- .env.example（API Key 模板，注意不是 .env 本身！）\n"
    "- .gitignore（忽略规则）\n"
    "- app.py（Web 服务器）\n"
    "- schematic_engine.py（AI 核心引擎）\n"
    "- templates/index.html（网页界面）\n"
    "- requirements.txt（依赖清单）\n"
    "- create_study_notes.py（笔记生成脚本）"
)

doc.add_heading("日常工作流程", level=3)
doc.add_paragraph(
    "以后每次修改代码后：\n\n"
    "1. git status          -- 看改了什么\n"
    "2. git add <file>      -- 选择要提交的文件\n"
    "3. git commit -m '...' -- 提交（写清楚改了什么）\n"
    "4. git push            -- 推送到 GitHub\n\n"
    "养成习惯：每完成一个小功能就 commit 一次！"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("阶段 3 完成！你的代码已经在 GitHub 上了！")
run.bold = True

doc.add_page_break()

# ===== 阶段 4 =====
doc.add_heading("阶段 4：CI/CD（持续集成/持续部署）", level=1)

doc.add_heading("Step 4.1：什么是 CI/CD？", level=2)
doc.add_paragraph(
    "CI/CD 是现代软件开发的核心实践。\n\n"
    "CI = Continuous Integration（持续集成）\n"
    "每次你 push 代码到 GitHub，系统自动运行测试。\n"
    "如果测试失败，你马上就知道哪里出了问题。\n\n"
    "CD = Continuous Deployment（持续部署）\n"
    "测试通过后，系统自动把代码部署到服务器上。\n"
    "不需要手动操作，推代码就自动上线。\n\n"
    "类比：\n"
    "- 没有 CI/CD：你写完代码 → 手动测试 → 手动上传服务器 → 祈祷不出bug\n"
    "- 有 CI/CD：你 push 代码 → 机器自动测试 → 自动部署 → 出问题马上通知你\n\n"
    "好处：\n"
    "- 减少人为错误（机器不会忘记跑测试）\n"
    "- 快速发现 bug（每次改动都测试）\n"
    "- 节省时间（不用手动部署）\n"
    "- 团队协作更安全（别人的代码也要通过测试才能合并）"
)

doc.add_heading("Step 4.2：什么是 GitHub Actions？", level=2)
doc.add_paragraph(
    "GitHub Actions 是 GitHub 内置的 CI/CD 工具，免费使用。\n\n"
    "核心概念：\n"
    "- Workflow（工作流）：一个自动化流程，定义在 .yml 文件里\n"
    "- Trigger（触发器）：什么时候运行？push 时？PR 时？定时？\n"
    "- Job（任务）：工作流里的一组步骤\n"
    "- Step（步骤）：一个具体的操作\n"
    "- Runner（运行器）：GitHub 免费提供的服务器，用来执行你的任务\n\n"
    "文件位置：\n"
    "  .github/workflows/ci.yml\n\n"
    "GitHub 会自动检测这个文件夹，发现 .yml 文件就会运行。"
)

doc.add_heading("Step 4.3：ci.yml 文件解读", level=2)
doc.add_paragraph(
    "我们的 CI 配置文件做了以下事情：\n\n"
    "1. 触发条件 (on:)\n"
    "   - push 到 master 分支时触发\n"
    "   - 有人提 Pull Request 到 master 时触发\n\n"
    "2. 运行环境 (runs-on:)\n"
    "   - ubuntu-latest：使用最新版 Ubuntu 服务器\n"
    "   - GitHub 免费提供，不需要你自己买服务器！\n\n"
    "3. 执行步骤 (steps:)\n"
    "   Step 1: Checkout code — 下载你的代码到服务器\n"
    "   Step 2: Setup Python — 安装 Python 3.11\n"
    "   Step 3: Install Graphviz — 安装画图工具\n"
    "   Step 4: Install dependencies — 安装 Python 库\n"
    "   Step 5: Run tests — 运行 pytest 测试！\n\n"
    "结果：\n"
    "- 全部通过 → 绿色勾 (success)\n"
    "- 有失败 → 红色 X (failure) → GitHub 会发邮件通知你"
)

doc.add_heading("Step 4.4：测试文件解读 (test_engine.py)", level=2)
doc.add_paragraph(
    "我们写了 10 个测试，覆盖了系统的各个部分：\n\n"
    "核心引擎测试：\n"
    "1. test_demo_mode_returns_data — 模拟模式能正常返回数据\n"
    "2. test_result_has_required_fields — 数据包含必要字段\n"
    "3. test_device_format — 设备格式正确（有 id, name, type）\n"
    "4. test_connection_format — 连接格式正确（有 from, to, cable）\n"
    "5. test_connections_reference_valid_devices — 连接指向存在的设备\n"
    "6. test_generate_diagram_creates_file — 能生成 PNG 文件\n"
    "7. test_system_prompt_quality — Prompt 包含关键指令\n\n"
    "Web 接口测试：\n"
    "8. test_flask_app_starts — 网站能启动\n"
    "9. test_generate_endpoint — /generate 接口返回正确数据\n"
    "10. test_empty_input_rejected — 空输入被正确拒绝"
)

doc.add_heading("Step 4.5：什么是 pytest？", level=2)
doc.add_paragraph(
    "pytest 是 Python 最流行的测试框架。\n\n"
    "规则很简单：\n"
    "- 测试文件以 test_ 开头\n"
    "- 测试函数以 test_ 开头\n"
    "- 用 assert 语句验证结果\n\n"
    "例如：\n"
    "  def test_addition():\n"
    "      assert 1 + 1 == 2      # 通过\n"
    "      assert 1 + 1 == 3      # 失败！\n\n"
    "运行命令：\n"
    "  python -m pytest test_engine.py -v\n\n"
    "  -v 表示 verbose（详细模式），显示每个测试的结果"
)

doc.add_heading("Step 4.6：查看 CI 结果", level=2)
doc.add_paragraph(
    "方法 1：GitHub 网页\n"
    "打开 https://github.com/yeahhao1997/ai-schematic-generator/actions\n"
    "可以看到每次 push 的测试结果\n\n"
    "方法 2：GitHub CLI\n"
    "在终端运行：gh run list --limit 5\n"
    "显示最近 5 次运行的状态\n\n"
    "我们的第一次 CI 运行结果：success！\n"
    "10 个测试全部通过，耗时约 30 秒。"
)

doc.add_heading("CI/CD 工作流程图", level=3)
doc.add_paragraph(
    "你 push 代码\n"
    "    |\n"
    "    v\n"
    "GitHub 检测到 push 事件\n"
    "    |\n"
    "    v\n"
    "启动 Ubuntu 服务器（Runner）\n"
    "    |\n"
    "    v\n"
    "下载你的代码\n"
    "    |\n"
    "    v\n"
    "安装 Python + 依赖\n"
    "    |\n"
    "    v\n"
    "运行 pytest 测试\n"
    "    |\n"
    "    v\n"
    "全部通过？ → 绿勾\n"
    "有失败？   → 红X + 邮件通知"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("阶段 4 完成！你的项目现在有了自动化 CI/CD！")
run.bold = True

doc.add_page_break()

# ===== 阶段 5 =====
doc.add_heading("阶段 5：部署上线", level=1)

doc.add_heading("Step 5.1：什么是部署？", level=2)
doc.add_paragraph(
    "部署 (Deploy) = 把你的程序放到互联网上，让任何人都能访问。\n\n"
    "之前你只能在自己电脑上运行（localhost），别人访问不了。\n"
    "部署后，你的系统有了一个公开网址，全世界都能用。\n\n"
    "类比：\n"
    "- localhost = 在自己家做饭，只有自己能吃\n"
    "- 部署 = 开了一家餐厅，任何人都能来吃"
)

doc.add_heading("Step 5.2：部署平台选择", level=2)
doc.add_paragraph(
    "常见的免费部署平台：\n\n"
    "PythonAnywhere（我们用的）\n"
    "- 优点：免费、不需要信用卡、对 Python/Flask 支持好\n"
    "- 缺点：免费版有流量限制\n"
    "- 网址：pythonanywhere.com\n\n"
    "Render.com\n"
    "- 优点：支持多种语言、自动部署\n"
    "- 缺点：现在需要信用卡\n\n"
    "其他选项：Railway、Fly.io、Vercel（适合前端项目）"
)

doc.add_heading("Step 5.3：PythonAnywhere 部署步骤", level=2)
doc.add_paragraph(
    "1. 注册 PythonAnywhere 免费账号\n"
    "   pythonanywhere.com → Create Beginner account\n\n"
    "2. 打开 Bash Console，运行：\n"
    "   git clone https://github.com/yeahhao1997/ai-schematic-generator.git\n"
    "   cd ai-schematic-generator\n"
    "   pip install --user -r requirements.txt\n"
    "   echo 'ANTHROPIC_API_KEY=your-key' > .env\n\n"
    "3. 创建 Web App\n"
    "   Web 菜单 → Add new web app → Flask → Python 3.13\n\n"
    "4. 修改 Source code 路径\n"
    "   改成：/home/yeahhao1997/ai-schematic-generator\n\n"
    "5. 修改 WSGI 配置文件\n"
    "   把默认内容替换为：\n"
    "   import sys, os\n"
    "   path = '/home/yeahhao1997/ai-schematic-generator'\n"
    "   if path not in sys.path:\n"
    "       sys.path.append(path)\n"
    "   os.chdir(path)\n"
    "   from dotenv import load_dotenv\n"
    "   load_dotenv(os.path.join(path, '.env'))\n"
    "   from app import app as application\n\n"
    "6. 点 Reload 按钮\n\n"
    "7. 访问 https://yeahhao1997.pythonanywhere.com"
)

doc.add_heading("Step 5.4：关键概念", level=2)
deploy_concepts = [
    ("WSGI (Web Server Gateway Interface)",
     "Python Web 应用和服务器之间的标准接口。PythonAnywhere 通过 WSGI 文件找到你的 Flask app。"),
    ("gunicorn",
     "生产级别的 Python Web 服务器。开发时用 Flask 自带的服务器，上线时用 gunicorn（更稳定、更快）。"),
    ("环境变量 (Environment Variables)",
     "服务器上的配置信息。API Key 等敏感信息通过环境变量传递，不写在代码里。"),
    ("render.yaml",
     "Render.com 的配置文件，定义如何构建和启动你的应用。其他平台有类似的配置方式。"),
]
for concept, explanation in deploy_concepts:
    p = doc.add_paragraph()
    r = p.add_run(f"{concept}：")
    r.bold = True
    p.add_run(f"\n{explanation}")

doc.add_heading("Step 5.5：更新部署的代码", level=2)
doc.add_paragraph(
    "以后你修改了代码，需要更新线上版本：\n\n"
    "1. 本地修改代码\n"
    "2. git add + git commit + git push（推到 GitHub）\n"
    "3. GitHub Actions 自动运行测试\n"
    "4. 去 PythonAnywhere 的 Bash Console 运行：\n"
    "   cd ~/ai-schematic-generator\n"
    "   git pull\n"
    "5. 去 Web 页面点 Reload\n\n"
    "完整流程：\n"
    "改代码 → push → CI 自动测试 → 手动 git pull → Reload → 上线"
)

doc.add_page_break()

# ===== 总结 =====
doc.add_heading("项目总结", level=1)
doc.add_paragraph(
    "恭喜！你从零完成了一个完整的 AI 项目！\n\n"
    "你学会了：\n"
    "1. Python 编程基础（函数、模块、JSON）\n"
    "2. 调用 AI API（Claude API + System Prompt）\n"
    "3. 自动画图（Graphviz）\n"
    "4. Web 开发（Flask + HTML/CSS/JavaScript）\n"
    "5. 版本控制（Git + GitHub）\n"
    "6. CI/CD（GitHub Actions + pytest）\n"
    "7. 部署上线（PythonAnywhere）\n\n"
    "你的系统：\n"
    "- 代码仓库：https://github.com/yeahhao1997/ai-schematic-generator\n"
    "- 线上网址：https://yeahhao1997.pythonanywhere.com\n"
    "- CI/CD：https://github.com/yeahhao1997/ai-schematic-generator/actions\n\n"
    "下一步可以做：\n"
    "- 拿到 Anthropic API Key，切换到真实 AI 模式\n"
    "- 在 Fiverr 开始接单\n"
    "- 添加更多功能（PDF 导出、draw.io 格式、更多图表类型）\n"
    "- 学习 Docker 容器化部署"
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("全部 5 个阶段完成！")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0, 102, 204)

# ===== 保存 =====
output_path = r"C:\Users\User\ai-schematic-generator\Study_Notes_AI_Schematic.docx"
doc.save(output_path)
print("Done!")
